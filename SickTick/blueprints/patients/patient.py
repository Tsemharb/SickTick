from os import listdir
from datetime import datetime
import time
import docx
import re
import json
# import random

# TODO: check dates by regexp (sometimes additional info presents)
# TODO: deal with get_correct_timestamp function variables outside of function

class Patient_parser:
    def __init__(self, id):

        self.tables_mapping = {'ФИО': 'general_info',
                          'Общие анализы крови': 'cbc',
                          'Биохимические анализы крови': 'blood_tests',
                          'Общие анализы мочи': 'urine_tests',
                          'Иммунологические исследования': 'immunologic_tests',
                          'Дополнительные исследования': 'additional_tests',
                          'Микробиологические посевы': 'microbiology',
                          'Этиотропная терапия': 'antibiotics',
                          'Температурный лист': 'temperature'}

# read antibiotics codes from json
        with open('AB_codes.json', 'rb') as json_file:
            self.json_codes = json.load(json_file)

        self.patient_file_name = self.get_patient_file_name(id)
        self.doc = docx.Document('data/' + self.patient_file_name)

        self.tables = {}
        for table in self.doc.tables:  #get tables from .doc file ('table_name': table_obj)
            try:
                self.tables[self.tables_mapping[table.rows[0].cells[0].text.strip()]] = table
            except:
                continue

        self.current_date = None
        self.date_counter = 1
        self.errors = []
        self.is_error = False

        self.general_info = self.get_general_info()
        print('got general info')

        try:
            self.temperature = self.get_temperature(self.tables['temperature'], 
                                                    self.general_info['admission_date'][-4:],
                                                    self.general_info['discharge_date'][-4:])
        except:
            self.temperature = []
        print('got temperature info')

        try:
            self.antibiotics = self.get_antibiotics(self.tables['antibiotics'])
        except:
            self.antibiotics = []
        print('got antibiotics info')

        self.additional_tests = self.get_additional_tests(self.tables['additional_tests'])
        print('got additional tests')

        try:
            self.cbc = self.get_tests(self.tables['cbc'], self.general_info['admission_date'][-4:], self.general_info['discharge_date'][-4:])
        except:
            self.is_error = True
            self.errors.append('Something is wrong with complete blood count')
        print('got complete blood count')

    @staticmethod
    def get_timestamp(date_str):
        a = int(datetime.timestamp(datetime.strptime(date_str, '%d.%m.%Y')) * 1000)
        return a

    @staticmethod
    def get_patient_file_name(id):
        files = [f.replace('.docx', '') for f in listdir('data')]
        for file in files:
            if file.split(' ')[0] == id:
                return file + '.docx'

    @staticmethod
    def get_font_size_from_text_length(length):
        in_min = 16  # minimum input length
        in_max = 219  # maximum input length
        out_min = 10  # minimum font size
        out_max = 19  #maximum font size
        return int(out_max - (length - in_min) * (out_max - out_min) / (in_max - in_min))

    @staticmethod
    def check_date_string(date):
        date = date.strip()
        m = re.search('(0|1|2|3)\d\.(0|1)\d\.(19|20)\d{2}', date)
        if m:
            return m.group(0)
        elif re.match(r'(0|1|2|3)\d\.(0|1)\d\.\d{2}', date):
            return date[:6] + '20' + date[-2:]
        return False


# method sets different timestamps for date duplicates (same dates won't have same x-coords)
    # @staticmethod
    def get_correct_timestamp(self, date, count):
        date_shift = 0;
        # print(self.current_date)
        if date == self.current_date:
            date_shift = self.date_counter / count * 24 * 60 * 60 * 1000
            self.date_counter+=1
        else:
            self.current_date = date
            self.date_counter = 1
            date_shift = 0
        return self.get_timestamp(date) + date_shift



    def get_general_info(self):
        try:
            gen_info_mapping = {'ФИО': 'name',
                                'Дата рождения': 'birthdate',
                                'Адрес домашний': 'address',
                                'Отделение': 'department',
                                'Дата госпитализации': 'admission_date',
                                'Дата выписки': 'discharge_date',
                                'Клинический диагноз': 'diagnosis',
                                'Диагноз заключительный': 'diagnosis'}
            general_info = {}
            the_table = self.tables['general_info']
            for row in the_table.rows:
                general_info[gen_info_mapping[row.cells[0].text.strip()]] = row.cells[1].text
            general_info['admission_date'] = self.check_date_string(general_info['admission_date'])
            general_info['discharge_date'] = self.check_date_string(general_info['discharge_date'])
            print(general_info['admission_date'])
            print()
            print()
            general_info['admission_timestamp'] = self.get_timestamp(general_info['admission_date'])
            general_info['discharge_timestamp'] = self.get_timestamp(general_info['discharge_date'])
        except:
            self.is_error = True
            self.errors.append('Something is wrong with general info')
            return
        return general_info


    def get_temperature(self, table, admission_year, discharge_year):
        try:
            same_year = admission_year == discharge_year  # check for change of the year while in hospital
            month = int(self.general_info['admission_date'][3:5])
            year = admission_year
            temperature = {}
            rows_num = len(table.rows)
            cells_num = len(table.rows[1].cells)
            temp_id = 0

            for row in range(1, rows_num, 3):
                dates = [cell.text[:5] for cell in table.rows[row].cells]  # take first 5 symbols to avoid trailing stuff (eg dots, spaces etc)
                results = [cell.text for cell in table.rows[row+1].cells]
                for i in range (len(dates)):
                    if dates[i] != '':
                        if not same_year and month > int(dates[i][-2:]):
                            year = discharge_year
                            month = int(dates[i][-2:])
                        else:
                            month = int(dates[i][-2:])
                        temperature[temp_id] = {'date': dates[i] + '.' + year,
                                                'id': temp_id,
                                                'temp': results[i].replace(',', '.')}
                        temp_id+=1
                    else:
                        break

            # add count for each unique date
            counts = {}
            dates = [temperature[i]['date'] for i in range(len(temperature))]
            for date in set(dates):
                for i in range(len(temperature)):
                    counts[i] = dates.count(temperature[i]['date'])

            # add timestamp for every temperature measurement
            for i in range(len(temperature)):
                timestamp = self.get_correct_timestamp(temperature[i]['date'], counts[i])
                temperature[i]['timestamp'] = timestamp;
        except:
            self.is_error = True
            self.errors.append('Something is wrong with the temperature data')
            return
        return temperature


    def get_antibiotics(self, table):
        ab_colors = ['aqua', 'brown', 'cyan', 'darkred','deeppink',
                     'fuchsia', 'lightgrey', 'olive', 'peru', 'tan',
                     'mediumblue', 'palegreen', 'rebeccapurple']
        antibiotics = []
        try:
            rows_num = len(table.rows)
            cells_num = len(table.rows[1].cells)
            for row in range(1, rows_num):
                row_content = []
                # get info from each row
                for cell in range(cells_num):
                    row_content.append(table.rows[row].cells[cell].text)
                # row_content example: ['с 04.09.2018 по 05.09.2018', 'Цефепим 2,0 * 2р. в/в']
                ab_name = row_content[1]
                ab_name = re.split("(\d)+", ab_name)[0]
                ab_dose = row_content[1].replace(ab_name, '')
                # add antibiotics dates could be several runs for each ab
                dates = row_content[0].split(',')
                for date in dates:
                    if ab_name == '':
                        continue
                    antibiotic = {}
                    antibiotic['name'] = ab_name.strip()
                    antibiotic['dose'] = ab_dose
                    antibiotic['color'] = ab_colors[row-1]
                    d = date.strip().split(' ')
                    try:
                        antibiotic['dates'] = {'begin': d[1][:10], 'end': d[3][:10]}
                    except:
                        antibiotic['dates'] = {'begin': '0', 'end': '0'}
                    try:
                        antibiotic['timestamps'] = {'begin': self.get_timestamp(d[1][:10]), 'end': self.get_timestamp(d[3][:10])}
                    except:
                        antibiotic['timestamps'] = {'begin': 0, 'end': 0}
                    if not (self.check_date_string(antibiotic['dates']['begin'])
                            and self.check_date_string(antibiotic['dates']['end'])):
                        antibiotic['dates'] = {'begin': '0', 'end': '0'}
                        antibiotic['timestamps'] = {'begin': 0, 'end': 0}
                    antibiotic['draw'] = True
                    try:
                        antibiotic['abbrev'] = self.json_codes[ab_name.lower().strip()]['abbrev']
                    except:
                        antibiotic['abbrev'] = ab_name[:3].upper()
                    antibiotics.append(antibiotic)
        except:
            self.is_error = True
            self.errors.append('Something is wrong with antibiotics')
            return
        return antibiotics

    def get_additional_tests(self, table): # group by test name
        try:
            additional_tests = {}
            rows_num = len(table.rows)
            for r_num in range(1, rows_num):
                # get the result of additional test
                test_raw_name = table.rows[r_num].cells[1].text.split(':')[0]
                result = table.rows[r_num].cells[1].text.replace(test_raw_name + ':', '').strip()

                #get font size for result message
                if len(result) <= 15:
                    result_font_size = 20
                elif len(result) >= 220:
                    result_font_size = 9
                else:
                    result_font_size = self.get_font_size_from_text_length(len(result))

                # set test attributes
                test = {}
                test['id'] = 'additional-' + str(r_num)
                test['date'] = self.check_date_string(table.rows[r_num].cells[0].text)
                test['y'] = None
                test['dx'] = 60
                test['dy'] = - 60
                test['draw'] = False
                test['result_font_size'] = result_font_size
                test['result_bold'] = False
                test['result_color'] = 'grey' #'coral'
                test['title_font_size'] = 15
                test['title_bold'] = True
                test['title_color'] = 'grey' #'pink'
                if test['date']:
                    test['timestamp'] = self.get_timestamp(test['date'])
                    test['timestamp_init'] = self.get_timestamp(test['date'])
                else:
                    test['timestamp'] = self.get_timestamp(self.general_info['admission_date'])
                    test['timestamp_init'] = self.get_timestamp(self.general_info['admission_date'])
                test['result'] = result
                name = test_raw_name.replace(' Заключение', '').replace(' Результат', '')
                try:
                    additional_tests[name]
                except:
                    additional_tests[name] = []
                additional_tests[name].append(test)
        except:
            self.is_error = True
            self.errors.append('Something is wrong with additional tests')
            return
        return additional_tests


    # def get_tests(self, table, admission_year, discharge_year):
    #     same_year = admission_year == discharge_year
    #     month = int(self.general_info['admission_date'][3:5])
    #     year = admission_year
    #     units = {}
    #     data = {}
    #     referent_col_name =''
    #     rows_num = len(table.rows)
    #     cells_num = len(table.rows[1].cells)
    #     blank_rows_ind = []
    #     is_referent_col = False  #referent column flag ("ед.измер", "Норма")

    #     #get referent name ("ед.измер", "Норма")
    #     if table.rows[1].cells[1].text == 'ед.изм' or table.rows[1].cells[1].text == 'Норма':
    #         referent_col_name = table.rows[1].cells[1].text
    #         is_referent_col = True

    #     #get blank rows
    #     for i in range(rows_num):
    #         row = table.rows[i]
    #         if (row.cells[0].text == '' and row.cells[1].text == ''):
    #             blank_rows_ind.append(i)
    #     blank_rows_ind.append(len(table.rows))

    #     #get sample units
    #     if is_referent_col:
    #         for row in range(1, rows_num):
    #             unit_name = table.rows[row].cells[0].text
    #             if unit_name != '':
    #                 units[table.rows[row].cells[0].text] = table.rows[row].cells[1].text
    #             else:
    #                 continue

    #     #get samples results
    #     start_row = 2
    #     id = 0
    #     start_col = 2 if is_referent_col else 1
    #     for blank in blank_rows_ind:
    #         dates = []
    #         for cell in range (start_col, cells_num):
    #             date = table.rows[start_row - 1].cells[cell].text.strip()
    #         # add year to yearless date
    #             if date != '':
    #                 if not same_year and month > int(date[-2:]):
    #                     year = discharge_year
    #                     month = int(date[-2:])
    #                 else:
    #                     month = int(date[-2:])
    #             else:
    #                 continue
    #             date = date + '.' + year
    #             dates.append(date)

    #         for row in range (start_row, blank):
    #             observation_name = table.rows[row].cells[0].text.strip()
    #             observation_series = []
    #             if observation_name not in data:
    #                 data[observation_name] = []
                
    #             for cell in range(start_col, cells_num):
    #                 observation_result = table.rows[row].cells[cell].text.strip()
    #                 if observation_result != '':
    #                     date = (dates[cell - start_col])
    #                     data[observation_name].append({'date': date, 'result': observation_result, 'timestamp': self.get_timestamp(date)})
    #                 else: 
    #                     continue
    #         start_row = blank + 2
    #     return{'referent_col_name': referent_col_name, 'units': units, 'data': data}


    def get_tests(self, table, admission_year, discharge_year):
        same_year = admission_year == discharge_year
        month = int(self.general_info['admission_date'][3:5])
        year = admission_year
        units = {}
        data = {}
        blank_rows_ind = []
        referent_col_name = ''
        rows_num = len(table.rows)
        cells_num = len(table.rows[1].cells)
        is_referent_col = False  #referent column flag ("ед.измер", "Норма")

        # get referent name ("ед.измер", "Норма")
        if table.rows[1].cells[1].text == 'ед.изм' or table.rows[1].cells[1].text == 'Норма':
            referent_col_name = table.rows[1].cells[1].text
            is_referent_col = True

        # get sample units and blank rows
        for i in range(1, rows_num):        
            s = [cell.text for cell in table.rows[i].cells]
            # get sample units
            if s[0] != '':
                units[s[0]] = s[1] if is_referent_col else ''
                data[s[0]] = []
            # get blank rows
            elif s[1] == '':
                blank_rows_ind.append(i)
        blank_rows_ind.append(rows_num)

        #get samples results
        start_row = 2
        start_col = 1 if is_referent_col else 0
        for blank in blank_rows_ind:
        # get tests dates
            dates = [cell.text.strip() for cell in table.rows[start_row - 1].cells if cell.text.strip() != ''][start_col:]
        # get samples results
            for row_num in range(start_row, blank):
                row = [cell.text.strip() for cell in table.rows[row_num].cells]
                tests = []
                for i in range(len(dates)):
                    if row[i + start_col + 1] != '':
                        if not same_year and month > int(dates[i][-2:]):
                            year = discharge_year
                            month = int(dates[i][-2:])
                        else:
                            month = int(dates[i][-2:])
                        tests.append({'date': dates[i] + '.' + year, 'result': row[i + start_col + 1], 'timestamp': self.get_timestamp(dates[i] + '.' + year)})
                data[row[0]] = data[row[0]] + tests
            start_row = blank + 2
        return{'referent_col_name': referent_col_name, 'units': units, 'data': data}


